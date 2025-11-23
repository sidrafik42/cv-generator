from flask import Flask, render_template, request, send_file, redirect, url_for
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.colors import HexColor
from PIL import Image as PILImage
import io

# Handle imports that may cause issues with static analysis
try:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    WD_ALIGN_PARAGRAPH = None

try:
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    WD_STYLE_TYPE = None

# Get the root directory (parent of cv_generator)
root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Create necessary directories
directories = ['generated', 'static/uploads']
for directory in directories:
    dir_path = os.path.join(root_dir, directory)
    os.makedirs(dir_path, exist_ok=True)

# Initialize Flask with correct template and static folders
app = Flask(__name__, 
            template_folder=os.path.join(root_dir, 'templates'),
            static_folder=os.path.join(root_dir, 'static'))
app.config['UPLOAD_FOLDER'] = os.path.join(root_dir, 'static', 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_cv', methods=['POST'])
def generate_cv():
    # Get form data
    name = request.form.get('name', '')
    email = request.form.get('email', '')
    phone = request.form.get('phone', '')
    address = request.form.get('address', '')
    profile = request.form.get('profile', '')
    
    # Education
    education_institutes = request.form.getlist('education_institute[]')
    education_degrees = request.form.getlist('education_degree[]')
    education_years = request.form.getlist('education_year[]')
    
    # Experience
    experience_companies = request.form.getlist('experience_company[]')
    experience_positions = request.form.getlist('experience_position[]')
    experience_years = request.form.getlist('experience_year[]')
    experience_descriptions = request.form.getlist('experience_description[]')
    
    # Skills
    skills = request.form.getlist('skill[]')
    
    # Add new fields
    # Languages
    languages = request.form.getlist('language[]')
    proficiencies = request.form.getlist('proficiency[]')
    
    # Additional information
    certifications = request.form.get('certifications', '')
    interests = request.form.get('interests', '')
    references = request.form.get('references', '')
    
    # Photo handling
    photo_path = None
    if 'photo' in request.files:
        photo = request.files['photo']
        if photo and photo.filename != '':
            photo_path = os.path.join(app.config['UPLOAD_FOLDER'], photo.filename or '')
            photo.save(photo_path)
    
    # Generate files
    docx_filename = f"{name.replace(' ', '_') if name else 'cv'}_CV.docx"
    pdf_filename = f"{name.replace(' ', '_') if name else 'cv'}_CV.pdf"
    
    # Create Word document
    create_word_document(name, email, phone, address, profile, 
                        education_institutes, education_degrees, education_years,
                        experience_companies, experience_positions, experience_years, experience_descriptions,
                        skills, languages, proficiencies, certifications, interests, references,
                        photo_path, docx_filename)
    
    # Create PDF document
    create_pdf_document(name, email, phone, address, profile, 
                       education_institutes, education_degrees, education_years,
                       experience_companies, experience_positions, experience_years, experience_descriptions,
                       skills, languages, proficiencies, certifications, interests, references,
                       photo_path, pdf_filename)
    
    # Return to home with success message
    return redirect(url_for('success', docx_file=docx_filename, pdf_file=pdf_filename))

@app.route('/success')
def success():
    docx_file = request.args.get('docx_file')
    pdf_file = request.args.get('pdf_file')
    return render_template('success.html', docx_file=docx_file, pdf_file=pdf_file)

@app.route('/download/<filename>')
def download_file(filename):
    generated_path = os.path.join(root_dir, 'generated', filename)
    return send_file(generated_path, as_attachment=True)

def create_word_document(name, email, phone, address, profile,
                        education_institutes, education_degrees, education_years,
                        experience_companies, experience_positions, experience_years, experience_descriptions,
                        skills, languages, proficiencies, certifications, interests, references,
                        photo_path, filename):
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add styles
    from docx.shared import RGBColor
    
    # Create custom styles
    styles = doc.styles
    
    # Check if styles already exist, if not create them
    try:
        title_style = styles['CVTitle']
    except KeyError:
        title_style = styles.add_style('CVTitle', 1)  # 1 corresponds to PARAGRAPH
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    try:
        heading_style = styles['CVHeading']
    except KeyError:
        heading_style = styles.add_style('CVHeading', 1)  # 1 corresponds to PARAGRAPH
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(14)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
    
    try:
        subheading_style = styles['CVSubHeading']
    except KeyError:
        subheading_style = styles.add_style('CVSubHeading', 1)  # 1 corresponds to PARAGRAPH
        subheading_font = subheading_style.font
        subheading_font.name = 'Arial'
        subheading_font.size = Pt(12)
        subheading_font.bold = True
        subheading_font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    try:
        normal_style = styles['CVNormal']
    except KeyError:
        normal_style = styles.add_style('CVNormal', 1)  # 1 corresponds to PARAGRAPH
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
    
    # Add name as heading with modern styling
    title_para = doc.add_paragraph(name, style='CVTitle')
    title_para.alignment = 1  # 1 corresponds to CENTER
    
    # Add contact information in a modern layout
    contact_info = []
    if email:
        contact_info.append(f"📧 {email}")
    if phone:
        contact_info.append(f"📱 {phone}")
    if address:
        contact_info.append(f"📍 {address}")
    
    if contact_info:
        contact_para = doc.add_paragraph(" | ".join(contact_info), style='CVNormal')
        contact_para.alignment = 1  # 1 corresponds to CENTER
        doc.add_paragraph()  # Add spacing
    
    # Add photo if provided (positioned at top right)
    if photo_path and os.path.exists(photo_path):
        try:
            # Add a table to position photo next to name
            doc.add_paragraph()  # Add spacing
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            
            # Left cell for photo
            left_cell = table.cell(0, 0)
            left_cell.width = Inches(2)
            left_cell.paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.5))
            
            # Right cell for profile (if exists)
            right_cell = table.cell(0, 1)
            if profile:
                right_cell.paragraphs[0].text = profile
                right_cell.paragraphs[0].style = 'CVNormal'
            
            doc.add_paragraph()  # Add spacing after table
        except Exception as e:
            print(f"Error adding photo: {e}")
            # Fallback to original method if photo placement fails
            if profile:
                doc.add_paragraph('PROFESSIONAL PROFILE', style='CVHeading')
                doc.add_paragraph(profile, style='CVNormal')
                doc.add_paragraph()  # Add spacing
    else:
        # Add profile section with modern styling if no photo
        if profile:
            doc.add_paragraph('PROFESSIONAL PROFILE', style='CVHeading')
            doc.add_paragraph(profile, style='CVNormal')
            doc.add_paragraph()  # Add spacing
    
    # Add education section with improved formatting
    if education_institutes and any(education_institutes):
        doc.add_paragraph('EDUCATION', style='CVHeading')
        for i in range(len(education_institutes)):
            if education_institutes[i]:
                # Institution and degree on one line
                institution_text = f"{education_institutes[i]}"
                if i < len(education_degrees) and education_degrees[i]:
                    institution_text += f" - {education_degrees[i]}"
                
                institution_para = doc.add_paragraph(institution_text, style='CVSubHeading')
                
                # Year and details on next line
                year_text = ""
                if i < len(education_years) and education_years[i]:
                    year_text = education_years[i]
                
                if year_text:
                    year_para = doc.add_paragraph(year_text, style='CVNormal')
                    year_para.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()  # Add spacing
    
    # Add experience section with improved formatting
    if experience_companies and any(experience_companies):
        doc.add_paragraph('WORK EXPERIENCE', style='CVHeading')
        for i in range(len(experience_companies)):
            if experience_companies[i]:
                # Company and position on one line
                company_text = f"{experience_companies[i]}"
                if i < len(experience_positions) and experience_positions[i]:
                    company_text += f" - {experience_positions[i]}"
                
                company_para = doc.add_paragraph(company_text, style='CVSubHeading')
                
                # Year on next line
                year_text = ""
                if i < len(experience_years) and experience_years[i]:
                    year_text = experience_years[i]
                
                if year_text:
                    year_para = doc.add_paragraph(year_text, style='CVNormal')
                    year_para.paragraph_format.left_indent = Inches(0.25)
                
                # Description
                if i < len(experience_descriptions) and experience_descriptions[i]:
                    desc_para = doc.add_paragraph(experience_descriptions[i], style='CVNormal')
                    desc_para.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_paragraph()  # Add spacing
    
    # Add skills section with modern presentation
    if skills and any(skills):
        doc.add_paragraph('SKILLS', style='CVHeading')
        skills_str = " | ".join([skill for skill in skills if skill])
        doc.add_paragraph(skills_str, style='CVNormal')
        doc.add_paragraph()  # Add spacing
    
    # Add languages section
    if languages and any(languages):
        doc.add_paragraph('LANGUAGES', style='CVHeading')
        for i in range(len(languages)):
            if languages[i]:
                language_text = languages[i]
                if i < len(proficiencies) and proficiencies[i]:
                    language_text += f" - {proficiencies[i]}"
                doc.add_paragraph(language_text, style='CVNormal')
        doc.add_paragraph()  # Add spacing
    
    # Add certifications section
    if certifications:
        doc.add_paragraph('CERTIFICATIONS', style='CVHeading')
        doc.add_paragraph(certifications, style='CVNormal')
        doc.add_paragraph()  # Add spacing
    
    # Add interests section
    if interests:
        doc.add_paragraph('INTERESTS', style='CVHeading')
        doc.add_paragraph(interests, style='CVNormal')
        doc.add_paragraph()  # Add spacing
    
    # Add references section
    if references:
        doc.add_paragraph('REFERENCES', style='CVHeading')
        doc.add_paragraph(references, style='CVNormal')
    
    # Save document
    generated_dir = os.path.join(root_dir, 'generated')
    os.makedirs(generated_dir, exist_ok=True)
    doc.save(os.path.join(generated_dir, filename))

def create_pdf_document(name, email, phone, address, profile,
                       education_institutes, education_degrees, education_years,
                       experience_companies, experience_positions, experience_years, experience_descriptions,
                       skills, languages, proficiencies, certifications, interests, references,
                       photo_path, filename):
    # Create PDF with modern styling
    generated_dir = os.path.join(root_dir, 'generated')
    os.makedirs(generated_dir, exist_ok=True)
    pdf_path = os.path.join(generated_dir, filename)
    
    # Create document with modern styling
    doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    
    # Custom styles for modern look
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=26,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=HexColor('#2C3E50'),
        leading=30
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=10,
        textColor=HexColor('#3498DB'),
        leading=16
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=6,
        textColor=HexColor('#2C3E50'),
        leading=14,
        fontName='Helvetica-Bold'
    )
    
    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=20,
        alignment=TA_CENTER,
        leading=14
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=10,
        leading=14
    )
    
    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=normal_style,
        leftIndent=20,
        bulletFontSize=11,
        bulletIndent=10
    )
    
    story = []
    
    # Title
    story.append(Paragraph(name, title_style))
    
    # Contact information with icons
    contact_info = []
    if email:
        contact_info.append(f"📧 {email}")
    if phone:
        contact_info.append(f"📱 {phone}")
    if address:
        contact_info.append(f"📍 {address}")
    
    if contact_info:
        contact_para = Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(contact_info), contact_style)
        story.append(contact_para)
    
    # Profile with photo if provided
    if photo_path and os.path.exists(photo_path):
        try:
            # Create a table to position photo next to profile
            from reportlab.platypus import Table, TableStyle, Image as ReportLabImage
            from reportlab.lib import colors
            
            # Prepare table data with photo and profile
            table_data = []
            
            # Create photo image (resize to fit)
            try:
                photo_img = ReportLabImage(photo_path, width=1.5*inch, height=1.5*inch)
            except Exception:
                # If image loading fails, use a placeholder or skip photo
                photo_img = None
            
            # Create row with photo and profile
            row_data = []
            
            # Add photo cell (always add something, even if empty)
            if photo_img:
                row_data.append(photo_img)
            else:
                row_data.append(Paragraph("", normal_style))  # Empty cell if photo fails
            
            # Add profile cell
            if profile:
                row_data.append(Paragraph(profile, normal_style))
            else:
                row_data.append(Paragraph("", normal_style))  # Empty cell if no profile
            
            # Only create table if we have at least one non-empty cell
            if photo_img or profile:
                table_data.append(row_data)
                
                # Create table with 2 columns
                profile_table = Table(table_data, colWidths=[1.5*inch, 4.5*inch])
                profile_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),  # Photo column
                    ('ALIGN', (1, 0), (1, -1), 'LEFT'),   # Profile column
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                ]))
                
                story.append(profile_table)
                story.append(Spacer(1, 12))
            else:
                # Fallback if both photo and profile are missing
                if profile:
                    story.append(Paragraph("PROFESSIONAL PROFILE", heading_style))
                    story.append(Paragraph(profile, normal_style))
        except Exception as e:
            print(f"Error adding photo to PDF: {e}")
            # Fallback to original method if photo placement fails
            if profile:
                story.append(Paragraph("PROFESSIONAL PROFILE", heading_style))
                story.append(Paragraph(profile, normal_style))
    else:
        # Add profile section if no photo
        if profile:
            story.append(Paragraph("PROFESSIONAL PROFILE", heading_style))
            story.append(Paragraph(profile, normal_style))
    
    story.append(Spacer(1, 12))
    
    # Education
    if education_institutes and any(education_institutes):
        story.append(Paragraph("EDUCATION", heading_style))
        for i in range(len(education_institutes)):
            if education_institutes[i]:
                # Institution and degree
                edu_text = f"<font color='#2C3E50'><b>{education_institutes[i]}</b></font>"
                if i < len(education_degrees) and education_degrees[i]:
                    edu_text += f" - {education_degrees[i]}"
                story.append(Paragraph(edu_text, subheading_style))
                
                # Year
                if i < len(education_years) and education_years[i]:
                    story.append(Paragraph(education_years[i], normal_style))
        story.append(Spacer(1, 12))
    
    # Experience
    if experience_companies and any(experience_companies):
        story.append(Paragraph("WORK EXPERIENCE", heading_style))
        for i in range(len(experience_companies)):
            if experience_companies[i]:
                # Company and position
                exp_text = f"<font color='#2C3E50'><b>{experience_companies[i]}</b></font>"
                if i < len(experience_positions) and experience_positions[i]:
                    exp_text += f" - {experience_positions[i]}"
                story.append(Paragraph(exp_text, subheading_style))
                
                # Year
                if i < len(experience_years) and experience_years[i]:
                    story.append(Paragraph(experience_years[i], normal_style))
                
                # Description
                if i < len(experience_descriptions) and experience_descriptions[i]:
                    story.append(Paragraph(f"• {experience_descriptions[i]}", bullet_style))
        story.append(Spacer(1, 12))
    
    # Skills
    if skills and any(skills):
        story.append(Paragraph("SKILLS", heading_style))
        skills_text = " | ".join([skill for skill in skills if skill])
        story.append(Paragraph(skills_text, normal_style))
        story.append(Spacer(1, 12))
    
    # Languages
    if languages and any(languages):
        story.append(Paragraph("LANGUAGES", heading_style))
        for i in range(len(languages)):
            if languages[i]:
                language_text = languages[i]
                if i < len(proficiencies) and proficiencies[i]:
                    language_text += f" - {proficiencies[i]}"
                story.append(Paragraph(language_text, normal_style))
        story.append(Spacer(1, 12))
    
    # Certifications
    if certifications:
        story.append(Paragraph("CERTIFICATIONS", heading_style))
        story.append(Paragraph(certifications, normal_style))
        story.append(Spacer(1, 12))
    
    # Interests
    if interests:
        story.append(Paragraph("INTERESTS", heading_style))
        story.append(Paragraph(interests, normal_style))
        story.append(Spacer(1, 12))
    
    # References
    if references:
        story.append(Paragraph("REFERENCES", heading_style))
        story.append(Paragraph(references, normal_style))
    
    # Build PDF
    doc.build(story)

if __name__ == '__main__':
    # Use the PORT environment variable if provided (for Render), otherwise default to 5000
    port = int(os.environ.get('PORT', 5000))
    # Bind to all interfaces (0.0.0.0) for Render, localhost for local development
    host = '0.0.0.0' if os.environ.get('PORT') else '127.0.0.1'
    app.run(debug=False, host=host, port=port)
